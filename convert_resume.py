#!/usr/bin/env python3
import os
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import html2text
import re

# Get paths
base_path = Path(__file__).parent
html_file = base_path / "Resume_Sujal_Atkari.html"
pdf_output = base_path / "Resume_Sujal_Atkari.pdf"
docx_output = base_path / "Resume_Sujal_Atkari.docx"

print(f"Converting resume...")
print(f"Input: {html_file}")
print(f"PDF Output: {pdf_output}")
print(f"DOCX Output: {docx_output}")

# Convert HTML to PDF using ReportLab
print("\n[1/2] Creating PDF...")
try:
    # Use html2text to convert HTML to markdown, then format for PDF
    h = html2text.HTML2Text()
    h.ignore_links = False
    h.ignore_images = False
    
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Extract clean text from HTML
    import re
    text_content = re.sub('<[^<]+?>', '', html_content)
    text_content = re.sub(r'\s+', ' ', text_content).strip()
    
    # Create PDF using ReportLab
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    c = canvas.Canvas(str(pdf_output), pagesize=letter)
    width, height = letter
    
    # Set up styles
    c.setFont("Helvetica-Bold", 18)
    c.drawString(50, height - 50, "SUJAL SAHADEV ATKARI")
    
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 70, "Mumbai, Maharashtra | sujalatkari.22@gmail.com | +91 8530240622")
    c.drawString(50, height - 85, "GitHub: github.com/SujalAtkari | LinkedIn: linkedin.com/in/sujal-atkari")
    
    # Add a line
    c.setLineWidth(1)
    c.line(50, height - 100, width - 50, height - 100)
    
    # Define starting position for content
    y_position = height - 120
    line_height = 14
    
    # Add sections
    sections = [
        ("PROFESSIONAL SUMMARY", 
         "Full Stack Developer specializing in MERN Stack and Java, with hands-on experience through internships and real-world projects. "
         "Skilled in building responsive web applications, integrating APIs, and developing scalable backend systems. "
         "Strong academic consistency with improving performance across semesters and active participation in national-level hackathons."),
        ("TECHNICAL SKILLS",
         "Languages: Java, JavaScript, SQL, HTML5, CSS3\n"
         "Backend: Node.js, Express.js, REST APIs\n"
         "Frontend: React.js, Tailwind CSS\n"
         "Database: MongoDB, MySQL\n"
         "Tools: Git, GitHub, Postman, VS Code"),
    ]
    
    for section_title, section_content in sections:
        if y_position < 150:
            c.showPage()
            y_position = height - 50
        
        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, y_position, section_title)
        y_position -= line_height + 3
        
        c.setFont("Helvetica", 9)
        if '\n' in section_content:
            for line in section_content.split('\n'):
                c.drawString(60, y_position, line)
                y_position -= line_height
        else:
            # Word wrap for long text
            words = section_content.split()
            line = ""
            for word in words:
                if len(line) + len(word) < 85:
                    line += word + " "
                else:
                    c.drawString(60, y_position, line)
                    y_position -= line_height
                    line = word + " "
            if line:
                c.drawString(60, y_position, line)
                y_position -= line_height
        
        y_position -= line_height
    
    c.save()
    print(f"✓ PDF created successfully: {pdf_output}")
except Exception as e:
    print(f"✗ Error creating PDF: {e}")

# Convert HTML to DOCX
print("[2/2] Creating DOCX...")
try:
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Extract text content from HTML
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("SUJAL SAHADEV ATKARI")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add contact info
    contact = doc.add_paragraph()
    contact_run = contact.add_run("Mumbai, Maharashtra | sujalatkari.22@gmail.com | +91 8530240622")
    contact_run.font.size = Pt(10)
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    links = doc.add_paragraph()
    links_run = links.add_run("GitHub: github.com/SujalAtkari | LinkedIn: linkedin.com/in/sujal-atkari")
    links_run.font.size = Pt(10)
    links.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add horizontal line
    doc.add_paragraph("_" * 80)
    
    # Professional Summary
    summary_title = doc.add_heading("PROFESSIONAL SUMMARY", level=1)
    summary_title.runs[0].font.size = Pt(12)
    summary_text = doc.add_paragraph(
        "Full Stack Developer specializing in MERN Stack and Java, with hands-on experience through internships and real-world projects. "
        "Skilled in building responsive web applications, integrating APIs, and developing scalable backend systems. "
        "Strong academic consistency with improving performance across semesters and active participation in national-level hackathons."
    )
    summary_text.runs[0].font.size = Pt(10)
    
    # Technical Skills
    skills_title = doc.add_heading("TECHNICAL SKILLS", level=1)
    skills_title.runs[0].font.size = Pt(12)
    skills = [
        ("Languages:", "Java, JavaScript, SQL, HTML5, CSS3"),
        ("Backend:", "Node.js, Express.js, REST APIs"),
        ("Frontend:", "React.js, Tailwind CSS"),
        ("Database:", "MongoDB, MySQL"),
        ("Tools:", "Git, GitHub, Postman, VS Code"),
    ]
    for skill_type, skill_content in skills:
        p = doc.add_paragraph()
        p_run = p.add_run(skill_type + " ")
        p_run.bold = True
        p_run.font.size = Pt(10)
        p_run = p.add_run(skill_content)
        p_run.font.size = Pt(10)
    
    # Professional Experience
    exp_title = doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
    exp_title.runs[0].font.size = Pt(12)
    
    exp_job = doc.add_paragraph()
    job_run = exp_job.add_run("Web Development Intern")
    job_run.bold = True
    job_run.font.size = Pt(11)
    
    exp_company = doc.add_paragraph("Smart Matrix, 2025")
    exp_company.runs[0].font.italic = True
    exp_company.runs[0].font.size = Pt(10)
    
    exp_bullets = [
        "Worked on MERN stack-based web application in a real development environment",
        "Developed responsive UI components using React.js",
        "Integrated REST APIs and handled frontend-backend communication",
        "Collaborated using Git and followed structured development practices"
    ]
    for bullet in exp_bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    # Projects
    proj_title = doc.add_heading("PROJECTS", level=1)
    proj_title.runs[0].font.size = Pt(12)
    
    projects = [
        {
            "name": "NextKart - E-commerce Web Application",
            "tech": "MERN Stack (MongoDB, Express.js, React.js, Node.js)",
            "year": "2025",
            "bullets": [
                "Developed a full-stack e-commerce platform with product listing and cart features",
                "Implemented dynamic UI with React and managed API integration",
                "Built responsive design ensuring smooth user experience"
            ]
        },
        {
            "name": "QuizDual - Java Quiz Application",
            "tech": "Java",
            "year": "2024",
            "bullets": [
                "Designed and developed a quiz system with multiple categories and scoring logic",
                "Implemented modular structure for easy scalability and maintenance",
                "Handled user interaction and result evaluation efficiently"
            ]
        },
        {
            "name": "RTS Website - Frontend Project",
            "tech": "HTML, CSS, JavaScript",
            "year": "2024",
            "bullets": [
                "Built a responsive and user-friendly website interface",
                "Focused on clean UI design and performance optimization"
            ]
        }
    ]
    
    for project in projects:
        p = doc.add_paragraph()
        p_run = p.add_run(project["name"])
        p_run.bold = True
        p_run.font.size = Pt(11)
        
        tech_p = doc.add_paragraph(f"{project['tech']} • {project['year']}")
        tech_p.runs[0].font.italic = True
        tech_p.runs[0].font.size = Pt(10)
        
        for bullet in project["bullets"]:
            doc.add_paragraph(bullet, style='List Bullet')
    
    # Certifications
    cert_title = doc.add_heading("CERTIFICATIONS & ACHIEVEMENTS", level=1)
    cert_title.runs[0].font.size = Pt(12)
    achievements = [
        "Oracle Cloud Infrastructure AI Foundations Associate (2025)",
        "SQL for Beginners: Learn SQL using MySQL – Scaler (2025)",
        "Cleared Round 1 – Smart India Hackathon (SIH) 2025 and advanced to next round"
    ]
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    # Education
    edu_title = doc.add_heading("EDUCATION", level=1)
    edu_title.runs[0].font.size = Pt(12)
    
    edu_items = [
        ("B.E. in Information Technology", "Mumbai University", "Expected Graduation: 2027 | CGPI (Till Semester 5): 7.86"),
        ("HSC (12th)", "", "73.33%"),
        ("SSC (10th)", "", "83.60%")
    ]
    
    for degree, institution, details in edu_items:
        p = doc.add_paragraph()
        p_run = p.add_run(degree)
        p_run.bold = True
        p_run.font.size = Pt(11)
        
        if institution:
            inst_p = doc.add_paragraph(institution)
            inst_p.runs[0].font.size = Pt(10)
        
        details_p = doc.add_paragraph(details)
        details_p.runs[0].font.size = Pt(10)
    
    doc.save(str(docx_output))
    print(f"✓ DOCX created successfully: {docx_output}")
except Exception as e:
    print(f"✗ Error creating DOCX: {e}")

print("\n✓ Conversion complete!")
